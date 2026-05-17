#!/usr/bin/env python3
# -*-- coding: utf-8 -*-
"""配台システム使い方.docx を見出し・表・番号付き手順付きで再生成する。

既存 docx のスクリーンショット（word/media/*.png）を流用し、本文は整理版に差し替える。
"""
from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph

REPO = Path(__file__).resolve().parents[1]
SRC_DOCX = REPO / "配台システム使い方.docx"
OUT_DOCX = REPO / "配台システム使い方.docx"
OUT_POLISHED = REPO / "配台システム使い方（整理版）.docx"
OUT_TMP = REPO / "配台システム使い方.docx.new"
BACKUP = REPO / "配台システム使い方.docx.bak"
MEDIA_DIR = REPO / ".tmp_dispatch_manual_media"

# コンパクトレイアウト（A4・余白狭め・行間詰め）
MARGIN_LR = Cm(1.15)
MARGIN_TB = Cm(1.2)
BODY_PT = Pt(9.5)
IMAGE_WIDTH = Cm(17.4)  # 余白込みで用紙幅いっぱいに近い
SPACE_BODY_AFTER = Pt(1.5)
SPACE_LIST_AFTER = Pt(0.5)
SPACE_HEADING1_BEFORE = Pt(8)
SPACE_HEADING1_AFTER = Pt(2)
SPACE_HEADING2_BEFORE = Pt(5)
SPACE_HEADING2_AFTER = Pt(1.5)


def extract_media(docx_path: Path, dest: Path) -> None:
    dest.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(docx_path) as zf:
        for name in zf.namelist():
            if name.startswith("word/media/") and not name.endswith("/"):
                target = dest / Path(name).name
                target.write_bytes(zf.read(name))


def _set_east_asia_font(style, name: str = "Yu Gothic") -> None:
    style.font.name = name
    rpr = style._element.rPr
    if rpr is None:
        rpr = style._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), name)


def tight(p: Paragraph, *, before: Pt | None = None, after: Pt | None = None) -> Paragraph:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.line_spacing = 1.0
    pf.space_before = before if before is not None else Pt(0)
    pf.space_after = after if after is not None else SPACE_BODY_AFTER
    return p


def set_doc_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = MARGIN_TB
    section.bottom_margin = MARGIN_TB
    section.left_margin = MARGIN_LR
    section.right_margin = MARGIN_LR

    normal = doc.styles["Normal"]
    _set_east_asia_font(normal)
    normal.font.size = BODY_PT
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = SPACE_BODY_AFTER

    for level, before, after, size in (
        (1, SPACE_HEADING1_BEFORE, SPACE_HEADING1_AFTER, Pt(13)),
        (2, SPACE_HEADING2_BEFORE, SPACE_HEADING2_AFTER, Pt(11)),
        (3, Pt(4), Pt(1), Pt(10)),
    ):
        h = doc.styles[f"Heading {level}"]
        _set_east_asia_font(h)
        h.font.size = size
        h.paragraph_format.space_before = before
        h.paragraph_format.space_after = after
        h.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    title = doc.styles["Title"]
    _set_east_asia_font(title)
    title.font.size = Pt(16)
    title.paragraph_format.space_after = Pt(2)

    for style_name in ("List Bullet", "List Number"):
        if style_name in doc.styles:
            ls = doc.styles[style_name]
            _set_east_asia_font(ls)
            ls.font.size = BODY_PT
            ls.paragraph_format.space_before = Pt(0)
            ls.paragraph_format.space_after = SPACE_LIST_AFTER
            ls.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_note(doc: Document, text: str, kind: str = "info") -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = BODY_PT
    run.bold = True
    if kind == "warn":
        run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    else:
        run.font.color.rgb = RGBColor(0x00, 0x50, 0x80)
    p.paragraph_format.left_indent = Cm(0.3)
    tight(p, before=Pt(2), after=Pt(2))


def add_path_table(doc: Document, rows: list[tuple[str, str]]) -> Table:
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "工場"
    hdr[1].text = "出力フォルダ（UNC）"
    for i, (factory, path) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = factory
        table.rows[i].cells[1].text = path
    small = Pt(8.5)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                tight(p, after=Pt(0))
                for r in p.runs:
                    r.font.size = small
    tight(doc.add_paragraph(), before=Pt(1), after=Pt(2))
    return table


def add_image(doc: Document, media_dir: Path, filename: str, caption: str | None = None) -> None:
    path = media_dir / filename
    if not path.is_file():
        tight(doc.add_paragraph(f"（画像なし: {filename}）"))
        return
    doc.add_picture(str(path), width=IMAGE_WIDTH)
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tight(last, before=Pt(2), after=Pt(0))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(8)
        tight(cap, before=Pt(0), after=Pt(2))


def build(doc: Document, media_dir: Path) -> None:
    # タイトル
    t = doc.add_heading("配台システム 使い方", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("工程管理 AI デスクトップ — 現場オペレーション手順")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(11)
    meta = doc.add_paragraph("更新: 2026年5月17日")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(8)
    meta.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    tight(meta, after=Pt(4))

    doc.add_heading("全体の流れ", level=1)
    flow = doc.add_paragraph()
    flow.add_run(
        "[アラジン] 仮配台・計画/実績の出力\n"
        "    ↓\n"
        "[工程管理 AI] 段階1 → タスク入力の確認・編集 → 段階2\n"
        "    ↓\n"
        "[確認] 納期管理ビュー / 設備ガント / 配台計画手動修正\n"
        "    ↓\n"
        "[任意] 段階3（配台試行）→ Excel 確認 → [アラジン] 本番反映"
    ).font.name = "Consolas"
    tight(flow, after=Pt(2))

    add_note(
        doc,
        "マスタの前提: master.xlsm（国分は国分masters.xlsm 等）の設定は別資料で完了しているものとします。",
    )

    # --- 1. アラジン ---
    doc.add_heading("1. 事前準備（アラジン）", level=1)

    doc.add_heading("1.1 生産計画入力（仮配台）", level=2)
    for item in [
        "アラジンの「生産計画入力」で、配台日・数量を「仮」として事前入力しておく。",
        "仮数量はこの時点では自動配台では無視される（段階2 の割付根拠にはならない）。",
        "仮を入れておかないと、加工計画 DATA から行を吸い出せない。",
        "システム配台を本番とする場合、確定後はアラジン側の仮計画も実績に合わせて更新する（納期管理ビュー参照）。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.2 加工計画 DATA の出力", level=2)
    steps = [
        "アラジンで「工程別生産計画問い合わせ」を開く。",
        "倉庫 … 湖南工場 520201 / 国分工場 511101。",
        "その他の条件は社内手順どおり設定し、「問合せ」を実行する。",
        "抽出結果の列設定は必ず初期化してから出力する。",
        "「フォルダ」で出力先を指定し、「出力」を実行する。",
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(s, style="List Number")
    add_image(doc, media_dir, "image1.png", "工程別生産計画問い合わせ（設定例）")
    add_image(doc, media_dir, "image2.png", "問合せ結果（列設定の初期化を忘れない）")
    add_path_table(
        doc,
        [
            (
                "湖南工場",
                r"\\192.168.0.101\共有フォルダ\湖南工場\湖南共有\生産管理システム\管理システム\●DATA\生産計画問合せ",
            ),
            (
                "国分工場",
                r"\\192.168.0.101\共有フォルダ\国分工場\国分共有\●配台AIシステム\DATA\計画",
            ),
        ],
    )
    add_image(doc, media_dir, "image3.png", "出力画面（設定例）")

    doc.add_heading("1.3 加工実績明細 DATA の出力", level=2)
    for s in [
        "「加工実績明細問合せ」を開く。",
        "倉庫 … 計画と同じコード（湖南 520201 / 国分 511101）。",
        "加工日付 … できるだけ短い期間（1 日分の実績があれば足りることが多い）。",
        "列設定を初期化し、「問合せ」→「フォルダ」指定→「出力」。",
    ]:
        doc.add_paragraph(s, style="List Number")
    add_image(doc, media_dir, "image4.png", "加工実績明細問合せ（設定例）")
    add_image(doc, media_dir, "image5.png", "問合せ結果")
    add_path_table(
        doc,
        [
            (
                "湖南工場",
                r"\\192.168.0.101\共有フォルダ\湖南工場\湖南共有\002  加工G\●検査表作成\加工実績明細DATA" + "\\",
            ),
            (
                "国分工場",
                r"\\192.168.0.101\共有フォルダ\国分工場\国分共有\●配台AIシステム\DATA\実績",
            ),
        ],
    )
    add_image(doc, media_dir, "image6.png", "出力画面（設定例）")

    # --- 2. 段階1 ---
    doc.add_heading("2. 工程管理 AI：段階1", level=1)
    for s in [
        "PMD.exe（配布物）を起動する。",
        "「実行・ログ」タブを開き、必要なら「ログをクリア」する。",
        "「段階1 実行」を押す。",
        "正常終了すると「配台計画_タスク入力」タブへ自動遷移する。",
    ]:
        doc.add_paragraph(s, style="List Number")
    add_image(doc, media_dir, "image7.png", "実行・ログタブ")
    add_image(doc, media_dir, "image8.png", "段階1 実行ボタン")
    add_image(doc, media_dir, "image10.png", "段階1 完了後のタブ遷移")

    # --- 3. タスク入力 ---
    doc.add_heading("3. 配台計画_タスク入力", level=1)
    doc.add_paragraph(
        "上部「操作・ソース」… 読込・保存・段階2 実行。下部「データ表」… 配台対象行。"
        "アコーディオンを閉じると表の表示領域が広くなる。"
    )
    add_image(doc, media_dir, "image11.png", "操作・ソースのアコーディオン")

    doc.add_heading("3.1 フィルタ（加工内容）", level=2)
    doc.add_paragraph(
        "「加工内容」列のフィルタ（▽）で EC のみ、EC,検査 のみなどに絞り込める（Excel と同様）。",
        style="List Bullet",
    )
    add_image(doc, media_dir, "image13.png")
    add_image(doc, media_dir, "image14.png")
    add_image(doc, media_dir, "image17.png", "加工内容「EC」の例")
    add_image(doc, media_dir, "image18.png", "加工内容「EC,検査」の例")
    add_image(doc, media_dir, "image19.png", "フィルタ適用後（該当行のみ表示）")

    doc.add_heading("3.2 配台試行順番・配台不要・上書き列", level=2)
    for item in [
        "自動配台は配台試行順番の昇順で設備・担当を割り付ける。",
        "「依頼NO」付近をドラッグ＆ドロップ（または ↑↓）で優先順を変更できる。",
        "「配台不要」セルをダブルクリックすると配台対象外（yes 相当・赤背景）になる。",
        "「加工速度_上書き」「原反投入日_上書き」はダブルクリックで編集（アラジン本体は書き換えない）。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_image(doc, media_dir, "image20.png", "配台試行順番の入れ替え")
    add_image(doc, media_dir, "image21.png", "配台不要のマーキング")
    add_image(doc, media_dir, "image22.png")
    add_image(doc, media_dir, "image23.png", "加工速度_上書き")
    add_image(doc, media_dir, "image24.png", "原反投入日_上書き")

    doc.add_heading("3.3 保存と段階2", level=2)
    for s in [
        "「操作・ソース」を開き、表の編集を「保存」する。",
        "「当日は配台しない」にチェックが入っていることを確認する。",
        "「段階2 実行」を押す（未保存のままでは実行できない）。",
    ]:
        doc.add_paragraph(s, style="List Number")
    add_note(doc, "保存後、自動的に「実行・ログ」タブへ切り替わり段階2 が走る。", "warn")
    add_image(doc, media_dir, "image25.png", "保存")
    add_image(doc, media_dir, "image26.png", "段階2 実行")
    add_image(doc, media_dir, "image27.png", "段階2 実行中")
    doc.add_paragraph(
        "完了メッセージが表示される。Gemini API 利用時は数分以上かかることがある。",
        style="List Bullet",
    )
    add_image(doc, media_dir, "image28.png", "段階2 完了")

    # --- 4. 確認 ---
    doc.add_heading("4. 結果の確認", level=1)

    doc.add_heading("4.1 納期管理ビュー", level=2)
    for item in [
        "納期判定列で NG の行を確認する。",
        "NG の場合は「回答納期」を「加工終了日時」の次稼働日にアラジン上で設定する（要点）。",
        "アラジン計画とシステム配台（シス配台）を並べて比較できる。",
        "システム配台を本番とするなら、仮入力したアラジン計画も整合させる。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_image(doc, media_dir, "image29.png", "納期管理ビュー")
    add_image(doc, media_dir, "image30.png", "アラジンとシス配台の比較")

    doc.add_heading("4.2 設備ガント（グラフィック）", level=2)
    doc.add_paragraph(
        "設備別タイムラインで割付の重なり・空白・担当をグラフィカルに評価する。",
        style="List Bullet",
    )
    add_image(doc, media_dir, "image32.png", "設備ガント")

    doc.add_heading("4.3 配台計画手動修正", level=2)
    doc.add_paragraph(
        "「工程・機械×日」… 機械名ごとの日別合計加工量（設備負荷の確認）。",
        style="List Bullet",
    )
    add_image(doc, media_dir, "image33.png", "工程・機械×日")
    doc.add_paragraph(
        "「タスク×日付」… 依頼NOごとにいつ・何 m 配台されているか。ダブルクリック編集・ドラッグで日付間移動。",
        style="List Bullet",
    )
    add_image(doc, media_dir, "image34.png", "タスク×日付")

    doc.add_heading("4.4 段階3（インタラクティブ配台試行）", level=2)
    add_note(
        doc,
        "段階3 は勤怠（master の人員カレンダー）に縛られるため、当日指定の加工数を必ず満たせるとは限らない。"
        "残業・休出で余力を増やすか、「タスク×日付」で人の判断を反映してから実行すること。",
        "warn",
    )
    for item in [
        "加工量をドラッグ＆ドロップで移動し、人の考えを反映できる。",
        "試行後のセルは（段階3前）と（段階3後）の2行表示。（段階3後）は編集しない。",
        "人員不足・割当不可は shortages ダイアログで確認する。",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    add_image(doc, media_dir, "image35.png", "段階3 の操作例")
    add_image(doc, media_dir, "image36.png", "段階3 の注意（将来拡張のメモ）")

    doc.add_heading("4.5 段階3 実行後・アラジン転記", level=2)
    for s in [
        "「保存 (JSON+xlsx)」で手修正を確定する（運用では保存を推奨）。",
        "「実行・ログ」の「エクセルを開く」で production_plan / member_schedule 等を確認する。",
        "納期管理ビューでも同系の数量・納期を確認できる。",
        "32インチ程度のワイドモニターでアラジンと本アプリを上下または左右に並べ、シス配台結果を転記する。",
    ]:
        doc.add_paragraph(s, style="List Number")
    add_image(doc, media_dir, "image37.png")
    add_image(doc, media_dir, "image38.png", "エクセルを開く")
    add_image(doc, media_dir, "image39.png", "アラジンへの転記（画面配置例）")

    # チェックリスト
    doc.add_heading("1日の締めチェックリスト", level=1)
    checks = [
        "段階1・段階2 がログ上正常終了している",
        "タスク入力表を保存済み",
        "納期管理ビューで NG 行に手順どおり対応した",
        "ガント・手動修正表で設備・人の偏りがない",
        "段階3 を使った場合、不足ダイアログを解消または記録した",
        "アラジン本番計画をシス配台と整合した",
    ]
    for c in checks:
        p = doc.add_paragraph(c, style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.35)
        tight(p, after=SPACE_LIST_AFTER)


def main() -> None:
    if not SRC_DOCX.is_file():
        raise SystemExit(f"not found: {SRC_DOCX}")
    if BACKUP.is_file():
        extract_media(BACKUP, MEDIA_DIR)
    else:
        shutil.copy2(SRC_DOCX, BACKUP)
        extract_media(BACKUP, MEDIA_DIR)
    doc = Document()
    set_doc_defaults(doc)
    build(doc, MEDIA_DIR)
    targets = [OUT_POLISHED, OUT_TMP]
    for target in targets:
        doc.save(target)
        print(f"wrote: {target}")
    try:
        OUT_TMP.replace(OUT_DOCX)
        print(f"updated: {OUT_DOCX}")
    except OSError:
        pass
    print(f"backup: {BACKUP}")


if __name__ == "__main__":
    main()
